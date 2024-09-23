import os
import json
import requests
import streamlit as st
import pandas as pd
import csv
#import pdfplumber
#import ollama
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain_core.language_models.llms import LLM
from langchain.globals import set_llm_cache
from langchain_community.vectorstores import Chroma
from langchain_community.cache import InMemoryCache
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from typing import Any, List, Mapping, Optional
from io import BytesIO
from datetime import datetime

# 초기 세션 상태 설정
if "messages" not in st.session_state:
    st.session_state["messages"] = [{"role": "assistant", "content": "무엇을 도와 드릴까요?"}]
if "vectorstore_cache" not in st.session_state:
    st.session_state["vectorstore_cache"] = {}
if "loading_text" not in st.session_state:
    st.session_state["loading_text"] = None
if "retriever" not in st.session_state:
    st.session_state["retriever"] = None
if 'questions' not in st.session_state:
    st.session_state['questions'] = []
if 'answers' not in st.session_state:
    st.session_state['answers'] = []
if "cache_initialized" not in st.session_state:
    set_llm_cache(InMemoryCache())
    st.session_state.cache_initialized = True
if 'chat_history' not in st.session_state:
    st.session_state['chat_history'] = []

# +

# HCX API 클래스
class LlmClovaStudio(LLM):
    """
    Custom LLM class for using the ClovaStudio API.
    """
    host: str
    api_key: str
    api_key_primary_val: str
    request_id: str

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.host = kwargs.get('host')
        self.api_key = kwargs.get('api_key')
        self.api_key_primary_val = kwargs.get('api_key_primary_val')
        self.request_id = kwargs.get('request_id')

    @property
    def _llm_type(self) -> str:
        return "custom"

    def _call(self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None
    ) -> str:
        """
        Make an API call to the ClovaStudio endpoint using the specified 
        prompt and return the response.
        """

        if stop is not None:
            raise ValueError("stop kwargs are not permitted.")

        headers = {
            "X-NCP-CLOVASTUDIO-API-KEY": self.api_key,
            "X-NCP-APIGW-API-KEY": self.api_key_primary_val,
            "X-NCP-CLOVASTUDIO-REQUEST-ID": self.request_id,
            "Content-Type": "application/json; charset=utf-8",
            "Accept": "text/event-stream"
        }

        sys_prompt = """
        당신은 AK아이에스의 사내 규정, 정책, 복리후생, 업무 가이드 등에 대해 답변하는 역할을 맡고 있습니다
        - Context 내에서만 답하며, 추측하거나 추가 정보 제공 금지
        - 감사, 칭찬에 대한 답변은 '감사합니다!'로만 응답하세요. 추가 설명 금지
        - Context에 정보가 없으면 '잘 모르겠습니다'라고 답변
        """
        #- 질문에 대해 간결하고 명확하게 핵심 정보만 전달, 공감 표현은 가능
        #- 회사, 업무, 복지(선물, 휴가 등)과 관련 없는 질문에는 "죄송합니다. 저는 업무 관련 내용에만 답변할 수 있습니다."라고 답변
        # 위 프롬프트 넣을 경우 과도하게 업무와 관련 없는 질문이라고 판단하는 경향 존재(개선 필요)
        
        preset_text = [{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt}]

        request_data = {
            "messages": preset_text,
            "topP": 0.6,
            "topK": 0,
            "maxTokens": 256,
            "temperature": 0.6,
            "repeatPenalty": 3,
            "stopBefore": [],
            "includeAiFilters": False
        }
        # API 요청
        response = requests.post(
            self.host + "/testapp/v1/chat-completions/HCX-003",  # 본인이 finetunning 한 API 경로 , 일반 HCX03써도 무관합니다
            headers=headers,
            json=request_data,
            stream=True
        )

        # 스트림에서 마지막 'data:' 라인을 찾기 위한 로직
        last_data_content = ""

        for line in response.iter_lines():
            if line:
                decoded_line = line.decode("utf-8")
                if '"data":"[DONE]"' in decoded_line:
                    break
                if decoded_line.startswith("data:"):
                    last_data_content = json.loads(decoded_line[5:])["message"]["content"]

        return last_data_content

llm = LlmClovaStudio(
    host='https://clovastudio.stream.ntruss.com',
    api_key='NTA0MjU2MWZlZTcxNDJiY6+5UNhJXWh3gqmFLbiMpde7ehpEJAFPwFUIey9lGc0S',
    api_key_primary_val='VhkfehtF14qpXmZPIA6VRw6x1c1eDCXp3P6BfbrG',
    #request_id='1e8ac996-b6e9-45f9-a64f-a64e37a029cf' #HCX-DASH-001
    request_id='b9288b57-8e12-45cc-b378-49cc13d8dbb6' #HCX-003
)

# -

#'Manual' 경로 내 pdf 파일 로딩
@st.cache_resource(show_spinner=False)
def extract_text_from_pdfs(folder_path, start_page=None, end_page=None):
    text = ''
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                start_page = start_page - 1 if start_page else 0
                end_page = end_page if end_page and end_page <= len(reader.pages) else len(reader.pages)
                for page_num in range(start_page, end_page):
                    text += reader.pages[page_num].extract_text()

    return text

# @st.cache_resource(show_spinner=False)
# def extract_text_from_pdfs(folder_path, start_page=None, end_page=None):
#     all_text = ''
    
#     for filename in os.listdir(folder_path):
#         if filename.endswith(".pdf"):
#             pdf_path = os.path.join(folder_path, filename)
#             try:
#                 with pdfplumber.open(pdf_path) as pdf:
#                     num_pages = len(pdf.pages)
                    
#                     # 시작 페이지와 끝 페이지 처리
#                     start = max(start_page - 1, 0) if start_page else 0
#                     end = min(end_page, num_pages) if end_page else num_pages
                    
#                     # 페이지별 텍스트 추출
#                     for page_num in range(start, end):
#                         page = pdf.pages[page_num]
#                         page_text = page.extract_text()
#                         if page_text:  # 텍스트가 있을 경우에만 추가
#                             all_text += page_text
#             except Exception as e:
#                 print(f"Error processing {filename}: {e}")
    
#     return all_text

# 'Manual' 경로 내 docx 파일 로딩
@st.cache_resource(show_spinner=False)
def extract_text_from_docx(folder_path):
    all_text = ''
    
    # 폴더 내 파일들을 확인
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            docx_path = os.path.join(folder_path, filename)
            try:
                # docx 파일 열기
                doc = DocxDocument(docx_path)
                
                # 문서의 모든 단락을 읽어와서 텍스트로 변환
                for para in doc.paragraphs:
                    all_text += para.text + '\n'
            except Exception as e:
                print(f"Error processing {filename}: {e}")
    
    return all_text

# 벡터DB 임베딩 및 retriver 생성
@st.cache_resource(show_spinner=False)
def retrieve_docs(text, model_index=0):
    vectorstore_path = 'vectorstore_' + str(model_index)
 
    model_list = [
        'bespin-global/klue-sroberta-base-continue-learning-by-mnr',
        'BAAI/bge-m3',
        'All-MiniLM-L6-v2',
        'sentence-transformers/paraphrase-MiniLM-L6-v2'
    ]

    embeddings = HuggingFaceEmbeddings(
        model_name=model_list[model_index],
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True},
    )

    os.makedirs(vectorstore_path, exist_ok=True)

    if os.path.exists(os.path.join(vectorstore_path, "index")):
        vectorstore = Chroma(persist_directory=vectorstore_path, embedding_function=embeddings)
    else:
        docs = [Document(page_content=text)]
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=300)
        splits = text_splitter.split_documents(docs)
        vectorstore = Chroma.from_documents(splits, embeddings, persist_directory=vectorstore_path)
        vectorstore.persist()

    return vectorstore.as_retriever(
        search_type='mmr',
        search_kwargs={'k': 3, 'lambda_mult': 0.5}
        )

# 줄바꿈 formatting
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

# RAG 핵심 함수
def rag_chain(question):
    question_history = "\n".join(str(x) for x in st.session_state['questions'][:-4:-1])
    retrieved_docs = retriever.invoke("\n".join([question, question_history])) # retriever 통한 벡터DB에서 필요한 정보 검색
    formatted_context = format_docs(retrieved_docs) # formatting
    formatted_prompt = f"Question: {question}\n\nContext: {formatted_context}" # API 전달할 Input 생성
    response = llm.invoke(formatted_prompt) # API 통해 답변 생성
    save_chat_to_xlsx(question, response, formatted_prompt) # 채팅 기록 저장
    return response

# -

# 채팅 기록을 CSV 파일로 저장하는 함수
def save_chat_to_csv(question, response, formatted_prompt):
    with open('chat_history.csv', mode='a', newline='', encoding='utf-8') as file:
    #with open('chat_history.csv', mode='a', newline='') as file:    
        writer = csv.writer(file)
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        writer.writerow([timestamp, question, response, formatted_prompt])

def save_chat_to_xlsx(question, response, formatted_prompt):
    file_path = 'chat_history.xlsx'
    
    # 새로운 데이터 행
    new_data = {
        'Timestamp': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        'Question': [question],
        'Response': [response],
        'Formatted Prompt': [formatted_prompt]
    }
    
    # 새로운 데이터를 DataFrame으로 변환
    new_df = pd.DataFrame(new_data)
    
    # 파일이 이미 있는지 확인하여 처리
    if os.path.exists(file_path):
        # 기존 Excel 파일에 데이터 추가
        existing_df = pd.read_excel(file_path)
        combined_df = pd.concat([existing_df, new_df], ignore_index=True)
    else:
        # 새로운 Excel 파일 생성
        combined_df = new_df

    combined_df.to_excel(file_path, index=False)

# -

if __name__ == "__main__":
    st.title("📟챗봇 아키입니다📟")
    st.markdown("AK아이에스의 규정, 복지제도에 대해 대답할 수 있어요.")
    st.markdown("현재 챗봇은 **테스트용**으로 운영되고 있습니다. 아껴주세요.")
    with st.spinner('불러오는 중입니다. 잠시만 기다려주세요.'):
        loading_text = extract_text_from_docx("Manual") # 매뉴얼 로딩
        loading_text = loading_text + extract_text_from_pdfs("Manual")
        retriever = retrieve_docs(loading_text) # retriever 생성

    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])

    if prompt := st.chat_input():
        prompt=prompt[:100]
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state['questions'].append(prompt)
        st.chat_message("user").write(prompt)
        with st.spinner(''):
            msg = rag_chain(prompt)
        st.session_state.messages.append({"role": "assistant", "content": msg})
        st.chat_message("assistant").write(msg)
        
#terminal 실행 커맨드
#streamlit run app.py